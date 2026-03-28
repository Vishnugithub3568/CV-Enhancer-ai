from docx import Document

# Create a simple test resume
doc = Document()
doc.add_heading('John Doe', 0)
doc.add_heading('Contact', 1)
doc.add_paragraph('Email: john@example.com')
doc.add_paragraph('Phone: (555) 123-4567')
doc.add_heading('Education', 1)
doc.add_paragraph('Bachelor of Science in Computer Science from State University, 2020')
doc.add_heading('Skills', 1)
doc.add_paragraph('Python, FastAPI, JavaScript, React, SQL, Docker')
doc.add_heading('Projects', 1)
doc.add_paragraph('CV Enhancer AI - Full stack application for resume parsing and enhancement')
doc.add_heading('Experience', 1)
doc.add_paragraph('Junior Developer at Tech Corp (2020-2022): Developed web applications using Python and FastAPI')

# Save to uploads folder
test_path = 'uploads/test_resume.docx'
doc.save(test_path)
print(f'Test resume created at: {test_path}')
