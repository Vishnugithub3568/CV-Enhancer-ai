import os
from typing import Optional

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.storage import GENERATED_DIR, ensure_directories

ensure_directories()


SECTION_ORDER = [
    ("summary", "Professional Summary"),
    ("education", "Education"),
    ("skills", "Technical Skills"),
    ("projects", "Projects"),
    ("experience", "Internship / Work Experience"),
    ("activities", "Achievements / Leadership / Activities"),
    ("certifications", "Certifications"),
]


def _safe_text(value) -> str:
    if isinstance(value, str):
        return value.strip()
    return ""


def _build_contact_line(data: dict) -> str:
    parts = []
    for key in ("email", "phone"):
        value = _safe_text(data.get(key, ""))
        if value:
            parts.append(value)

    return " | ".join(parts)


def _build_links_lines(data: dict) -> list:
    links = data.get("links", {}) if isinstance(data.get("links"), dict) else {}
    lines = []

    for key in ("linkedin", "github", "portfolio"):
        value = _safe_text(links.get(key, ""))
        if value:
            lines.append(value)

    other_links = links.get("other", []) if isinstance(links.get("other"), list) else []
    for item in other_links:
        value = _safe_text(item)
        if value:
            lines.append(value)

    return lines


def _parse_section_lines(value: str) -> list:
    lines = []
    for raw in value.split("\n"):
        line = raw.strip()
        if not line:
            continue
        is_bullet = line.startswith("- ") or line.startswith("* ")
        text = line[2:].strip() if is_bullet else line
        lines.append((is_bullet, text))
    return lines


def _merge_activity_sections(data: dict) -> str:
    merged = []
    for key in ("achievements", "responsibilities", "extracurricular"):
        value = _safe_text(data.get(key, ""))
        if not value:
            continue
        for is_bullet, text in _parse_section_lines(value):
            merged.append(f"- {text}" if is_bullet else f"- {text}")
    return "\n".join(merged).strip()


def _get_section_value(data: dict, key: str) -> str:
    if key == "activities":
        return _merge_activity_sections(data)
    return _safe_text(data.get(key, ""))


def _add_docx_section_content(doc: Document, value: str) -> None:
    parsed_lines = _parse_section_lines(value)
    if not parsed_lines:
        doc.add_paragraph(value)
        return

    for is_bullet, text in parsed_lines:
        if is_bullet:
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)


def _escape_pdf_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _append_pdf_section_content(story: list, value: str, body_style: ParagraphStyle, bullet_style: ParagraphStyle) -> None:
    parsed_lines = _parse_section_lines(value)
    if not parsed_lines:
        story.append(Paragraph(_escape_pdf_text(value), body_style))
        return

    for is_bullet, text in parsed_lines:
        safe_text = _escape_pdf_text(text)
        if is_bullet:
            story.append(Paragraph(f"- {safe_text}", bullet_style))
        else:
            story.append(Paragraph(safe_text, body_style))


def _create_docx(data: dict, file_path: str) -> None:
    doc = Document()

    name = _safe_text(data.get("name", "")) or "Professional Resume"
    doc.add_heading(name, level=0)

    contact = _build_contact_line(data)
    if contact:
        doc.add_paragraph(contact)

    for key, title in SECTION_ORDER:
        value = _get_section_value(data, key)
        if not value:
            continue
        doc.add_heading(title, level=1)
        _add_docx_section_content(doc, value)

    links_lines = _build_links_lines(data)
    if links_lines:
        doc.add_heading("Coding Profiles / Links", level=1)
        for link in links_lines:
            doc.add_paragraph(link, style="List Bullet")

    doc.save(file_path)


def _create_pdf(data: dict, file_path: str) -> None:
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#123A7A"),
        spaceBefore=10,
        spaceAfter=4,
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        leading=14,
        spaceAfter=6,
    )

    bullet_style = ParagraphStyle(
        "BulletBody",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=0,
        spaceAfter=4,
    )

    doc = SimpleDocTemplate(file_path, pagesize=LETTER, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    story = []

    name = _safe_text(data.get("name", "")) or "Professional Resume"
    story.append(Paragraph(name, styles["Title"]))

    contact = _build_contact_line(data)
    if contact:
        story.append(Paragraph(contact, body_style))

    story.append(Spacer(1, 8))

    for key, title in SECTION_ORDER:
        value = _get_section_value(data, key)
        if not value:
            continue
        story.append(Paragraph(title, heading_style))
        _append_pdf_section_content(story, value, body_style, bullet_style)

    links_lines = _build_links_lines(data)
    if links_lines:
        story.append(Paragraph("Coding Profiles / Links", heading_style))
        for link in links_lines:
            story.append(Paragraph(f"- {_escape_pdf_text(link)}", bullet_style))

    doc.build(story)


def generate_resume(enhanced_data: dict, base_filename: str = "generated_resume", include_docx: bool = True) -> dict:
    pdf_path = os.path.join(GENERATED_DIR, f"{base_filename}.pdf")
    _create_pdf(enhanced_data, pdf_path)

    docx_path: Optional[str] = None
    if include_docx:
        docx_path = os.path.join(GENERATED_DIR, f"{base_filename}.docx")
        _create_docx(enhanced_data, docx_path)

    return {
        "pdf_path": pdf_path,
        "docx_path": docx_path,
    }
